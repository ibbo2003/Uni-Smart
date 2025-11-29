# app.py
from flask import Flask, request, jsonify
from flask_cors import CORS
import mysql.connector
import os
from dotenv import load_dotenv
from seating_algorithm import arrange_seats
from auth_middleware import require_auth, require_admin_or_faculty, require_admin
import traceback
import pandas as pd
from docx import Document
import PyPDF2
import re
import io

load_dotenv()
app = Flask(__name__)
CORS(app)

db_config = {
    'host': os.getenv('DB_HOST'),
    'user': os.getenv('DB_USER'),
    'password': os.getenv('DB_PASSWORD'),
    'database': 'unismart_db'
}

# ================== EXAM ROOMS ENDPOINTS ==================
@app.route('/rooms', methods=['GET'])
@require_auth()  # Any authenticated user can view rooms
def get_all_rooms():
    """Get all exam rooms - Any authenticated user"""

    conn = mysql.connector.connect(**db_config)
    try:
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT * FROM exam_rooms ORDER BY id")
        rooms = cursor.fetchall()
        return jsonify(rooms), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if conn and conn.is_connected():
            cursor.close()
            conn.close()

@app.route('/rooms', methods=['POST'])
@require_admin_or_faculty
def create_room():
    """Create a new exam room - ADMIN and FACULTY only"""

    data = request.json
    room_id = data.get('id')
    num_rows = data.get('num_rows')
    num_cols = data.get('num_cols')

    if not all([room_id, num_rows, num_cols]):
        return jsonify({"error": "Missing required fields"}), 400

    capacity = num_rows * num_cols

    conn = mysql.connector.connect(**db_config)
    try:
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO exam_rooms (id, num_rows, num_cols, capacity) VALUES (%s, %s, %s, %s)",
            (room_id, num_rows, num_cols, capacity)
        )
        conn.commit()
        return jsonify({"message": "Room created successfully", "id": room_id}), 201
    except mysql.connector.IntegrityError:
        return jsonify({"error": "Room ID already exists"}), 409
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if conn and conn.is_connected():
            cursor.close()
            conn.close()

@app.route('/rooms/<room_id>', methods=['PUT'])

@require_admin_or_faculty
def update_room(room_id):
    """Update an existing exam room - ADMIN and FACULTY only"""

    data = request.json
    num_rows = data.get('num_rows')
    num_cols = data.get('num_cols')

    if not all([num_rows, num_cols]):
        return jsonify({"error": "Missing required fields"}), 400

    capacity = num_rows * num_cols

    conn = mysql.connector.connect(**db_config)
    try:
        cursor = conn.cursor()
        cursor.execute(
            "UPDATE exam_rooms SET num_rows = %s, num_cols = %s, capacity = %s WHERE id = %s",
            (num_rows, num_cols, capacity, room_id)
        )
        conn.commit()

        if cursor.rowcount == 0:
            return jsonify({"error": "Room not found"}), 404

        return jsonify({"message": "Room updated successfully"}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if conn and conn.is_connected():
            cursor.close()
            conn.close()

@app.route('/rooms/<room_id>', methods=['DELETE'])

@require_admin_or_faculty
def delete_room(room_id):
    """Delete an exam room - ADMIN and FACULTY only"""

    conn = mysql.connector.connect(**db_config)
    try:
        cursor = conn.cursor()
        cursor.execute("DELETE FROM exam_rooms WHERE id = %s", (room_id,))
        conn.commit()

        if cursor.rowcount == 0:
            return jsonify({"error": "Room not found"}), 404

        return jsonify({"message": "Room deleted successfully"}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if conn and conn.is_connected():
            cursor.close()
            conn.close()

# ================== EXAMS ENDPOINTS ==================
@app.route('/exams', methods=['GET'])

@require_auth()  # Any authenticated user can view exams
def get_all_exams():
    """Get all exams - Any authenticated user"""

    conn = mysql.connector.connect(**db_config)
    try:
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT * FROM exams ORDER BY exam_date, exam_session")
        exams = cursor.fetchall()
        return jsonify(exams), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if conn and conn.is_connected():
            cursor.close()
            conn.close()

@app.route('/exams', methods=['POST'])

@require_admin_or_faculty
def create_exam():
    """Create a new exam - ADMIN and FACULTY only"""

    data = request.json
    subject_code = data.get('subject_code')
    exam_date = data.get('exam_date')
    exam_session = data.get('exam_session')

    if not all([subject_code, exam_date, exam_session]):
        return jsonify({"error": "Missing required fields"}), 400

    if exam_session not in ['morning', 'afternoon']:
        return jsonify({"error": "Invalid exam_session. Must be 'morning' or 'afternoon'"}), 400

    conn = mysql.connector.connect(**db_config)
    try:
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO exams (subject_code, exam_date, exam_session) VALUES (%s, %s, %s)",
            (subject_code, exam_date, exam_session)
        )
        conn.commit()
        exam_id = cursor.lastrowid
        return jsonify({"message": "Exam created successfully", "id": exam_id}), 201
    except mysql.connector.IntegrityError:
        return jsonify({"error": "Exam already exists for this subject, date, and session"}), 409
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if conn and conn.is_connected():
            cursor.close()
            conn.close()

@app.route('/exams/<int:exam_id>', methods=['DELETE'])

@require_admin_or_faculty
def delete_exam(exam_id):
    """Delete an exam - ADMIN and FACULTY only"""

    conn = mysql.connector.connect(**db_config)
    try:
        cursor = conn.cursor()
        cursor.execute("DELETE FROM exams WHERE id = %s", (exam_id,))
        conn.commit()

        if cursor.rowcount == 0:
            return jsonify({"error": "Exam not found"}), 404

        return jsonify({"message": "Exam deleted successfully"}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if conn and conn.is_connected():
            cursor.close()
            conn.close()

# ================== EXAM REGISTRATIONS ENDPOINTS ==================
@app.route('/registrations', methods=['POST'])
@require_admin_or_faculty
def create_registration():
    """Register a student for an exam - ADMIN and FACULTY only"""
    data = request.json
    student_usn = data.get('student_usn')
    exam_id = data.get('exam_id')

    if not all([student_usn, exam_id]):
        return jsonify({"error": "Missing required fields"}), 400

    conn = mysql.connector.connect(**db_config)
    try:
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO exam_registrations (student_usn, exam_id) VALUES (%s, %s)",
            (student_usn, exam_id)
        )
        conn.commit()
        registration_id = cursor.lastrowid
        return jsonify({"message": "Registration created successfully", "id": registration_id}), 201
    except mysql.connector.IntegrityError as e:
        if 'Duplicate entry' in str(e):
            return jsonify({"error": "Student already registered for this exam"}), 409
        return jsonify({"error": "Foreign key constraint failed. Check student_usn and exam_id"}), 400
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if conn and conn.is_connected():
            cursor.close()
            conn.close()

def extract_usns_from_excel(file):
    """Extract USNs from Excel file"""
    try:
        df = pd.read_excel(file, engine='openpyxl')
        # Look for USN column (case insensitive)
        usn_column = None
        for col in df.columns:
            if 'usn' in str(col).lower():
                usn_column = col
                break

        if usn_column is None:
            # Try first column if no USN column found
            usn_column = df.columns[0]

        usns = df[usn_column].dropna().astype(str).str.strip().tolist()
        return [usn for usn in usns if usn and usn.upper() != 'USN']
    except Exception as e:
        print(f"Error extracting from Excel: {e}")
        return []

def extract_usns_from_word(file):
    """Extract USNs from Word document"""
    try:
        doc = Document(file)
        usns = []
        usn_pattern = re.compile(r'\b[A-Z0-9]{10}\b')  # VTU USN pattern

        for paragraph in doc.paragraphs:
            matches = usn_pattern.findall(paragraph.text)
            usns.extend(matches)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    matches = usn_pattern.findall(cell.text)
                    usns.extend(matches)

        return list(set(usns))  # Remove duplicates
    except Exception as e:
        print(f"Error extracting from Word: {e}")
        return []

def extract_usns_from_pdf(file):
    """Extract USNs from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        usns = []
        usn_pattern = re.compile(r'\b[A-Z0-9]{10}\b')  # VTU USN pattern

        for page in pdf_reader.pages:
            text = page.extract_text()
            matches = usn_pattern.findall(text)
            usns.extend(matches)

        return list(set(usns))  # Remove duplicates
    except Exception as e:
        print(f"Error extracting from PDF: {e}")
        return []

@app.route('/registrations/batch', methods=['POST'])
@require_admin_or_faculty
def create_batch_registrations():
    """Register multiple students for an exam - ADMIN and FACULTY only"""
    print("\n" + "="*60)
    print("[BATCH REGISTRATION] Received request")
    print("="*60)

    student_usns = []
    exam_id = None

    # Check if it's a file upload
    if 'file' in request.files:
        file = request.files['file']
        exam_id = request.form.get('exam_id')

        print(f"[FILE UPLOAD] Filename: {file.filename}")
        print(f"[FILE UPLOAD] Exam ID: {exam_id}")

        if not file or not exam_id:
            return jsonify({"error": "Missing file or exam_id"}), 400

        filename = file.filename.lower()

        # Extract USNs based on file type
        if filename.endswith(('.xlsx', '.xls')):
            print("[PROCESSING] Excel file detected")
            student_usns = extract_usns_from_excel(file)
        elif filename.endswith('.docx'):
            print("[PROCESSING] Word file detected")
            student_usns = extract_usns_from_word(file)
        elif filename.endswith('.pdf'):
            print("[PROCESSING] PDF file detected")
            student_usns = extract_usns_from_pdf(file)
        else:
            return jsonify({"error": "Unsupported file format. Use Excel, Word, or PDF"}), 400

    # JSON input (original method)
    elif request.json:
        data = request.json
        exam_id = data.get('exam_id')
        student_usns = data.get('student_usns', [])
        print("[JSON INPUT] Received JSON data")
        print(f"[JSON INPUT] Exam ID: {exam_id}")
        print(f"[JSON INPUT] USNs count: {len(student_usns)}")

    else:
        return jsonify({"error": "No data provided. Send JSON or upload a file"}), 400

    if not exam_id:
        return jsonify({"error": "exam_id is required"}), 400

    if not student_usns:
        return jsonify({"error": "No valid student USNs found"}), 400

    print(f"\n[VALIDATION] Total USNs extracted: {len(student_usns)}")
    print(f"[VALIDATION] Sample USNs: {student_usns[:5]}")

    conn = mysql.connector.connect(**db_config)
    try:
        cursor = conn.cursor()

        # Check if exam exists
        cursor.execute("SELECT id FROM exams WHERE id = %s", (exam_id,))
        if not cursor.fetchone():
            return jsonify({"error": f"Exam with ID {exam_id} not found"}), 404

        # Check which students exist in the database
        format_strings = ','.join(['%s'] * len(student_usns))
        cursor.execute(
            f"SELECT usn FROM students WHERE usn IN ({format_strings})",
            tuple(student_usns)
        )
        existing_students = {row[0] for row in cursor.fetchall()}

        print(f"[DATABASE CHECK] Found {len(existing_students)} existing students in database")

        if len(existing_students) == 0:
            return jsonify({
                "error": "None of the provided USNs exist in the students table",
                "provided_usns": student_usns[:10],
                "suggestion": "Please ensure students are added to the system first"
            }), 400

        # Filter to only existing students
        valid_usns = [usn for usn in student_usns if usn in existing_students]
        invalid_usns = [usn for usn in student_usns if usn not in existing_students]

        print(f"[FILTERING] Valid USNs: {len(valid_usns)}")
        print(f"[FILTERING] Invalid USNs: {len(invalid_usns)}")

        # Prepare batch insert for valid students only
        values = [(usn, exam_id) for usn in valid_usns]

        # Use INSERT IGNORE to skip duplicates
        cursor.executemany(
            "INSERT IGNORE INTO exam_registrations (student_usn, exam_id) VALUES (%s, %s)",
            values
        )
        conn.commit()

        registered_count = cursor.rowcount
        print(f"\n[SUCCESS] Registered {registered_count} students")
        print("="*60 + "\n")

        response = {
            "message": f"Successfully registered {registered_count} students",
            "registered_count": registered_count,
            "total_provided": len(student_usns),
            "valid_usns": len(valid_usns),
            "invalid_usns": len(invalid_usns)
        }

        if invalid_usns:
            response["invalid_usn_samples"] = invalid_usns[:5]

        return jsonify(response), 201

    except Exception as e:
        print(f"\n[ERROR] {str(e)}")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500
    finally:
        if conn and conn.is_connected():
            cursor.close()
            conn.close()

@app.route('/exams/<int:exam_id>/registrations', methods=['GET'])
@require_auth()  # Any authenticated user can view registrations
def get_exam_registrations(exam_id):
    """Get all students registered for a specific exam - Any authenticated user"""
    conn = mysql.connector.connect(**db_config)
    try:
        cursor = conn.cursor(dictionary=True)
        cursor.execute("""
            SELECT r.registration_id, r.student_usn, s.name, s.section_id
            FROM exam_registrations r
            JOIN students s ON r.student_usn = s.usn
            WHERE r.exam_id = %s
            ORDER BY r.student_usn
        """, (exam_id,))
        registrations = cursor.fetchall()
        return jsonify(registrations), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if conn and conn.is_connected():
            cursor.close()
            conn.close()

@app.route('/extract-students-from-pdf', methods=['POST'])
@require_admin_or_faculty
def extract_students_from_pdf():
    """Extract student data from uploaded PDF file - ADMIN and FACULTY only

    Intelligently extracts USN, Name, and Gender (optional) from PDF tables.
    Handles various column formats and orders.
    """
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400

    if not file.filename.endswith('.pdf'):
        return jsonify({"error": "File must be a PDF"}), 400

    try:
        import pdfplumber

        students = []

        def find_column_index(headers, keywords):
            """Find column index by matching keywords (case-insensitive)"""
            if not headers:
                return None

            for i, header in enumerate(headers):
                if header:
                    header_lower = str(header).lower().strip()
                    for keyword in keywords:
                        if keyword.lower() in header_lower:
                            return i
            return None

        def is_valid_usn(value):
            """Check if value looks like a valid VTU USN"""
            if not value:
                return False
            value = str(value).strip()
            # VTU USN format: typically 10 characters, alphanumeric
            # Must have both letters and digits
            if len(value) < 10:
                return False

            # Check if it's alphanumeric (allowing some special chars like .)
            clean_value = value.replace('.', '').replace('-', '').replace('_', '')
            if not clean_value:
                return False

            has_alpha = any(c.isalpha() for c in clean_value)
            has_digit = any(c.isdigit() for c in clean_value)

            # Should be mostly alphanumeric
            alphanumeric_count = sum(c.isalnum() for c in clean_value)

            return has_alpha and has_digit and alphanumeric_count >= 10

        def extract_usn_from_row(row):
            """Try to find USN in any column of the row"""
            for cell in row:
                if cell and is_valid_usn(cell):
                    return str(cell).strip()
            return None

        # Read PDF with pdfplumber for better table extraction
        with pdfplumber.open(file) as pdf:
            for page_num, page in enumerate(pdf.pages):
                print(f"\n[PDF EXTRACT] Processing page {page_num + 1}")

                # Extract tables from the page
                tables = page.extract_tables()

                if not tables:
                    print(f"[PDF EXTRACT] No tables found on page {page_num + 1}")
                    continue

                for table_num, table in enumerate(tables):
                    if not table or len(table) < 2:
                        continue

                    print(f"[PDF EXTRACT] Processing table {table_num + 1} with {len(table)} rows")

                    # Try to detect header row and column indices
                    header_row = table[0] if table else []
                    usn_col = find_column_index(header_row, ['usn', 'university', 'seat', 'number', 'no', 'roll'])
                    name_col = find_column_index(header_row, ['name', 'student'])
                    gender_col = find_column_index(header_row, ['gender', 'sex', 'g', 'male', 'female'])

                    print(f"[PDF EXTRACT] Detected columns - USN: {usn_col}, Name: {name_col}, Gender: {gender_col}")

                    # Process data rows (skip header)
                    for row_num, row in enumerate(table[1:], start=2):
                        if not row or len(row) < 2:
                            continue

                        # Extract USN - try detected column first, then search all columns
                        usn = None
                        if usn_col is not None and usn_col < len(row):
                            usn = row[usn_col]

                        # If USN column detection failed, search for USN pattern in row
                        if not is_valid_usn(usn):
                            usn = extract_usn_from_row(row)

                        if not is_valid_usn(usn):
                            continue  # Skip row if no valid USN found

                        # Extract Name - try detected column, then try adjacent columns
                        name = ""
                        if name_col is not None and name_col < len(row):
                            name = row[name_col] if row[name_col] else ""
                        else:
                            # Try to find name in columns near USN
                            for i, cell in enumerate(row):
                                if cell and str(cell).strip() and not is_valid_usn(cell):
                                    # Check if it looks like a name (mostly letters)
                                    cell_str = str(cell).strip()
                                    if len(cell_str) > 3 and sum(c.isalpha() or c.isspace() for c in cell_str) / len(cell_str) > 0.7:
                                        name = cell_str
                                        break

                        # Extract Gender (optional) - default to "Not Specified" if not found
                        gender = "Not Specified"
                        if gender_col is not None and gender_col < len(row) and row[gender_col]:
                            gender_value = str(row[gender_col]).strip().upper()
                            # Normalize gender values
                            if gender_value in ['M', 'MALE', 'BOY']:
                                gender = "Male"
                            elif gender_value in ['F', 'FEMALE', 'GIRL']:
                                gender = "Female"
                            else:
                                gender = gender_value if gender_value else "Not Specified"

                        # Clean up the data
                        usn = str(usn).strip().upper() if usn else ""
                        name = str(name).strip().upper() if name else "UNKNOWN"

                        if usn:  # Only add if we have a valid USN
                            student = {
                                "usn": usn,
                                "name": name,
                                "gender": gender
                            }
                            students.append(student)
                            print(f"[PDF EXTRACT] Row {row_num}: {student}")

        # Remove duplicates (same USN)
        unique_students = {}
        for student in students:
            if student['usn'] not in unique_students:
                unique_students[student['usn']] = student

        students = list(unique_students.values())

        if not students:
            return jsonify({"error": "No valid student data found in PDF. Please ensure the PDF contains a table with USN and Name columns."}), 400

        print(f"\n[PDF EXTRACT] Successfully extracted {len(students)} unique students")

        return jsonify({
            "message": f"Successfully extracted {len(students)} students",
            "students": students,
            "count": len(students)
        }), 200

    except Exception as e:
        print(f"[PDF EXTRACT ERROR] {str(e)}")
        traceback.print_exc()
        return jsonify({"error": f"Failed to extract data from PDF: {str(e)}"}), 500

@app.route('/students/batch-create', methods=['POST'])
@require_admin_or_faculty
def batch_create_students():
    """Create multiple students from extracted PDF data - ADMIN and FACULTY only"""
    data = request.json
    students = data.get('students', [])

    if not students:
        return jsonify({"error": "No student data provided"}), 400

    conn = mysql.connector.connect(**db_config)
    try:
        cursor = conn.cursor()

        # Prepare batch insert - using REPLACE to handle duplicates
        values = [(s['usn'], s['name'], s.get('section_id', '7_A')) for s in students]
        cursor.executemany(
            "INSERT INTO students (usn, name, section_id) VALUES (%s, %s, %s) "
            "ON DUPLICATE KEY UPDATE name = VALUES(name), section_id = VALUES(section_id)",
            values
        )
        conn.commit()

        return jsonify({
            "message": f"Successfully created/updated {len(students)} students",
            "count": len(students)
        }), 201

    except Exception as e:
        print(f"Error in batch create: {str(e)}")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500
    finally:
        if conn and conn.is_connected():
            cursor.close()
            conn.close()

# ================== SEATING PLAN ENDPOINTS ==================

@app.route('/generate_seating', methods=['POST'])
@require_admin_or_faculty
def generate_seating_plan():
    """Generate seating plan - ADMIN and FACULTY only"""
    user = request.current_user
    print('-----------------------------------------')
    print(f'[EXAM_SERVICE] Received request from {user["role"]} user (ID: {user["id"]})')
    data = request.json
    exam_date = data['exam_date']
    exam_session = data['exam_session']
    exam_type = data.get('exam_type', 'external')  # 'internal' or 'external'

    # DEBUG: Print received parameters
    print('[DEBUG] Received parameters:')
    print(f'  exam_date = {repr(exam_date)} (type: {type(exam_date).__name__})')
    print(f'  exam_session = {repr(exam_session)} (type: {type(exam_session).__name__})')
    print(f'  exam_type = {repr(exam_type)} (type: {type(exam_type).__name__})')

    conn = mysql.connector.connect(**db_config)
    try:
        cursor = conn.cursor(dictionary=True)

        print('[EXAM_SERVICE] [DB] Fetching data from database...')

        # 1. Fetch all available exam rooms
        cursor.execute("SELECT id, num_rows, num_cols FROM exam_rooms")
        rooms_data = cursor.fetchall()
        print(f"   - Found {len(rooms_data)} rooms.")

        # DEBUG: Show what exams exist in database
        cursor.execute("SELECT id, subject_code, exam_date, exam_session FROM exams")
        all_exams = cursor.fetchall()
        print(f'[DEBUG] Database contains {len(all_exams)} exams:')
        for exam in all_exams:
            print(f'  Exam ID {exam["id"]}: {exam["subject_code"]} on {exam["exam_date"]} {exam["exam_session"]}')

        # 2. Fetch all students registered for exams on the given date/session
        query = """
            SELECT r.student_usn, e.subject_code
            FROM exam_registrations r
            JOIN exams e ON r.exam_id = e.id
            WHERE e.exam_date = %s AND e.exam_session = %s
        """
        print(f'[DEBUG] Executing query with date={repr(exam_date)}, session={repr(exam_session)}')
        cursor.execute(query, (exam_date, exam_session))
        registrations = cursor.fetchall()

        print(f'[DEBUG] Query returned {len(registrations)} registrations')
        if len(registrations) > 0:
            print(f'[DEBUG] First 3 registrations: {registrations[:3]}')

        # Group students by subject
        students_by_subject = {}
        for reg in registrations:
            subject = reg['subject_code']
            if subject not in students_by_subject:
                students_by_subject[subject] = []
            students_by_subject[subject].append(reg['student_usn'])

        print(f"   - Found {len(registrations)} student registrations across {len(students_by_subject)} subjects.")

        # 3. Run the seating algorithm
        print(f'[EXAM_SERVICE] Starting seating arrangement algorithm for {exam_type} exam...')
        seating_plan = arrange_seats(rooms_data, students_by_subject, exam_type)

        print('[EXAM_SERVICE] Sending response back to the gateway...')
        return jsonify(seating_plan)

    except Exception as e:
        print(f"[EXAM_SERVICE] ERROR - An error occurred:")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500
    finally:
        if conn and conn.is_connected():
            cursor.close()
            conn.close()



@app.route('/seating-plan', methods=['GET'])
@require_auth()  # Any authenticated user can view seating plans
def get_seating_plan():
    """Get saved seating plan for a specific exam date and session - Any authenticated user"""

    exam_date = request.args.get('exam_date')
    exam_session = request.args.get('exam_session')

    if not exam_date or not exam_session:
        return jsonify({"error": "Missing exam_date or exam_session parameters"}), 400

    conn = mysql.connector.connect(**db_config)
    try:
        cursor = conn.cursor(dictionary=True)

        query = """
            SELECT
                esp.student_usn,
                s.name as student_name,
                e.subject_code,
                esp.room_id,
                esp.row_num,
                esp.col_num
            FROM exam_seating_plan esp
            JOIN students s ON esp.student_usn = s.usn
            JOIN exams e ON esp.exam_id = e.id
            WHERE e.exam_date = %s AND e.exam_session = %s
            ORDER BY esp.room_id, esp.row_num, esp.col_num
        """
        cursor.execute(query, (exam_date, exam_session))
        seating_plan = cursor.fetchall()

        return jsonify(seating_plan), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if conn and conn.is_connected():
            cursor.close()
            conn.close()


if __name__ == '__main__':
    # Run on a different port than the timetable service
    app.run(debug=True, port=5001)